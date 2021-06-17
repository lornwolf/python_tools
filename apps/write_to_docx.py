import json
import requests

from docx import Document


old_file = "C:/temp/input/农村房屋安全信息采集表（自建房）.docx"
new_path = "C:/temp/output/"
authorization = "Bearer eyJhbGciOiJIUzUxMiJ9.eyJzdWIiOiI4NjEzNjdkOC02OWNkLTRhMmEtYTA0Zi0xNzc5ZTgzNWQ1ODgiLCJleHAiOjE2MjU1OTAxODAsImlhdCI6MTYyMzg2MjE4MH0.AN7PAZ6VMVBZjbdJ55x7Bb7EpGXlrvXl9cSfSagNbNeaXaPG5A5p4XBaNCUnR0b7-Rcv5pGsT3T3ICN70KkZHA"
url = "https://ruralhouse.cubigdata.cn/houseInfoManage/queryHouseBaseInfoDetailById?id="


def check_and_change():
    """
    注意事项：
    直接操作paragraph.text会导致字体格式丢失，所以要遍历其run操作。
    一个段落会被分成多个run，为了确保替换元内容不被分割，最好用"AA、AB"这种形式，而不要使用类似"${21}"这样的字符串。
    """
    doc = Document(old_file)

    # 产权人
    AA = ""
    # 身份证号
    AB = ""
    # 建筑层数
    BA = "□"
    BB = "□"
    BC = "□"
    # 建筑面积
    BD = ""
    # 建造年代
    CA = "□"
    CB = "□"
    CC = "□"
    CD = "□"
    CE = "□"
    # 结构类型
    DA = "□"
    DB = "□"
    # 建造方式
    EA = "□"
    EB = "□"
    EC = "□"
    ED = "□"
    EE = ""
    # 土地性质
    FA = "□"
    FB = "□"
    #
    #
    # 加载事先取得的，包含所有人员信息的JSON文件。
    users = json.load(open('C:/temp/input/users.txt', 'r', encoding='utf-8'))
    print("共 " + str(len(users['data']['records'])) + " 名人员。")

    for user in users['data']['records']:
        response = requests.get(url + user['id'], headers={'Authorization': authorization})
        if response.status_code != 200:
            print('查询出错！')
            break
        detail = json.loads(response.text)
        # 产权人
        AA = detail["data"]["propertyOwner"]
        # 身份证号
        AB = detail["data"]["certNumber"]
        # 建筑层数
        buildStorey = detail["data"]["buildStorey"]
        if buildStorey == "1":
            BA = "■"
        elif buildStorey == "2":
            BB = "■"
        elif buildStorey == "3":
            BC = "■"
        # 建筑面积
        BD = detail["data"]["buildArea"]
        # 建造年代
        buildYear = detail["data"]["buildYear"]
        if buildYear == "1":
            CA = "■"
        elif buildYear == "2":
            CB = "■"
        elif buildYear == "3":
            CC = "■"
        elif buildYear == "4":
            CD = "■"
        elif buildYear == "5":
            CE = "■"
        # 结构类型
        structType = detail["data"]["structType"]
        if structType == "2": # 砖混结构（现浇楼板） / 砖石结构（非预制板）
            DA = "■"
        elif structType == "3": # 土木结构
            DB = "■"
        else:
            print("出现其他结构类型：" + AA)
            break
        # 建造方式
        buildMode = detail["data"]["buildMode"]
        buildModeOther = detail["data"]["buildModeOther"]
        if buildMode == "1":
            EA = "■"
        elif buildMode == "2":
            EB = "■"
        elif buildMode == "3":
            EC = "■"
        elif buildMode == "4":
            ED = "■"
            EE = buildModeOther
        # 土地性质
        landType = detail["data"]["landType"]
        if landType == "1":
            FA = "■"
        elif landType == "2":
            FB = "■"
        break

    # 遍历所有表格进行替换。
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        print(run.text)
                        run.text = run.text.replace("AA", AA)
                        run.text = run.text.replace("AB", AB)
                        run.text = run.text.replace("BA", BA)
                        run.text = run.text.replace("BB", BB)
                        run.text = run.text.replace("BC", BC)
                        run.text = run.text.replace("BD", str(BD))
                        run.text = run.text.replace("CA", CA)
                        run.text = run.text.replace("CB", CB)
                        run.text = run.text.replace("CC", CC)
                        run.text = run.text.replace("CD", CD)
                        run.text = run.text.replace("CE", CE)
                        run.text = run.text.replace("DA", DA)
                        run.text = run.text.replace("DB", DB)
                        run.text = run.text.replace("EA", EA)
                        run.text = run.text.replace("EB", EB)
                        run.text = run.text.replace("EC", EC)
                        run.text = run.text.replace("ED", ED)
                        run.text = run.text.replace("EE", EE)
                        run.text = run.text.replace("FA", FA)
                        run.text = run.text.replace("FB", FB)

    doc.save(new_path + AA + "（" + AB + "）.docx")


def main():
    check_and_change()


if __name__ == '__main__':
    main()
