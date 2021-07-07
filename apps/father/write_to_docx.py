import json
import os

from docx import Document

input_file = "C:/temp/input/农村房屋安全信息采集表（自建房）.docx"
file_path = "C:/temp/output/西村房屋信息/"
output_path = "C:/temp/output/农村房屋安全信息采集表/"

def check_and_change():
    """
    从住宅信息的JSON文件中读取信息，填入word模板文件中。
    注意事项：
    直接操作paragraph.text会导致字体格式丢失，所以要遍历其run操作。
    一个段落会被分成多个run，为了确保替换元内容不被分割，最好用"AA、AB"这种形式，而不要使用类似"${21}"这样的字符串。
    """

    for root, dirs, files in os.walk(file_path):

        # 遍历文件
        for file in files:
            # 产权人
            AA = ""
            # 身份证号
            AB = ""
            # 建筑层数
            BA = "□"
            BB = "□"
            BC = "□"
            # 建筑面积
            BD = ""
            # 建造年代
            CA = "□"
            CB = "□"
            CC = "□"
            CD = "□"
            CE = "□"
            # 结构类型
            DA = "□"
            DB = "□"
            DC = "□"
            DD = "□"
            DE = "□"
            # 建造方式
            EA = "□"
            EB = "□"
            EC = "□"
            ED = "□"
            EE = ""
            # 土地性质
            FA = "□"
            FB = "□"
            # 宅基地手续
            GA = "□"
            GB = "□"
            # 规划建设手续
            GC = "□"
            GD = "□"
            # 竣工验收手续
            HA = "□"
            HB = "□"
            # 房屋登记手续
            HC = "□"
            HD = "□"
            # 是否改造
            IA = "□"
            IB = "□"
            # 改造回数
            IC = "□"
            ID = "□"
            # 改造内容
            JA = "□"
            JB = "□"
            JC = "□"
            JD = "□"
            JE = "□"
            JF = ""
            # 是否用作经营
            KA = "□"
            KB = "□"
            # 经营审批手续
            LA = "□"
            LB = "□"
            # 主要用途
            MA = "□"
            MB = "□"
            MC = "□"
            # 采暖用能
            NA = "□"
            NB = "□"
            NC = "□"
            ND = "□"
            NE = "□"
            NF = ""
            # 炊事用能
            OA = "□"
            OB = "□"
            OC = "□"
            OD = "□"
            OE = "□"
            OF = ""
            doc = Document(input_file)
            detail = json.load(open(file_path + file, 'r', encoding='utf-8'))
            # 产权人
            AA = detail["data"]["propertyOwner"]
            print(AA)
            # 身份证号
            AB = detail["data"]["certNumber"]
            # 建筑层数
            buildStorey = detail["data"]["buildStorey"]
            if buildStorey == "1":
                BA = "■"
            elif buildStorey == "2":
                BB = "■"
            elif buildStorey == "3":
                BC = "■"
            # 建筑面积
            BD = detail["data"]["buildArea"]
            # 建造年代
            buildYear = detail["data"]["buildYear"]
            if buildYear == "1":
                CA = "■"
            elif buildYear == "2":
                CB = "■"
            elif buildYear == "3":
                CC = "■"
            elif buildYear == "4":
                CD = "■"
            elif buildYear == "5":
                CE = "■"
            # 结构类型
            structType = detail["data"]["structType"]
            if structType == "1":  # 砖混结构（预制楼板）
                DA = "■"
            elif structType == "2":  # 砖混结构（现浇楼板） / 砖石结构（非预制板）
                DB = "■"
            elif structType == "3":  # 土木结构
                DC = "■"
            elif structType == "6":  # 钢结构
                DD = "■"
            elif structType == "7":  # 土木结构
                DE = "■"
            else:
                print("出现其他结构类型：" + AA)
                break
            # 建造方式
            buildMode = detail["data"]["buildMode"]
            buildModeOther = detail["data"]["buildModeOther"]
            if buildMode == "1":
                EA = "■"
            elif buildMode == "2":
                EB = "■"
            elif buildMode == "3":
                EC = "■"
            elif buildMode == "4":
                ED = "■"
                EE = buildModeOther
            # 土地性质
            landType = detail["data"]["landType"]
            if landType == "1":
                FA = "■"
            elif landType == "2":
                FB = "■"
            # 宅基地手续
            homesteadPermit = detail["data"]["homesteadPermit"]
            if homesteadPermit == "1":
                GA = "■"
            elif homesteadPermit == "2":
                GB = "■"
            else:
                print("出现其他宅基地手续：" + AA)
                break
            # 规划建设手续
            planPermit = detail["data"]["planPermit"]
            if planPermit == "1":
                GC = "■"
            elif planPermit == "2":
                GD = "■"
            else:
                print("出现其他规划建设手续：" + AA)
                break
            # 竣工验收手续
            completePermit = detail["data"]["completePermit"]
            if completePermit == "1":
                HA = "■"
            elif completePermit == "2":
                HB = "■"
            else:
                print("出现其他竣工验收手续：" + AA)
                break
            # 房屋登记手续
            houseRegisterPermit = detail["data"]["houseRegisterPermit"]
            if houseRegisterPermit == "1":
                HC = "■"
            elif houseRegisterPermit == "2":
                HD = "■"
            else:
                print("出现其他竣工验收手续：" + AA)
                break
            # 是否改造
            houseTransformType = detail["data"]["houseTransformType"]
            if houseTransformType == "1":
                IB = "■"  # 是
                # 改造回数
                houseTransformNum = detail["data"]["houseTransformNum"]
                if houseTransformNum == "1":
                    IC = "■"
                elif houseTransformNum == "2":
                    ID = "■"
                else:
                    print("出现其他改造回数：" + AA)
                    break
                # 改造内容
                houseTransformContent = detail["data"]["houseTransformContent"]
                if len(houseTransformContent) > 1:
                    print("出现多次改造内容：" + AA)
                    break
                if houseTransformContent[0] == "1":
                    JA = "■"
                elif houseTransformContent[0] == "2":
                    JB = "■"
                else:
                    print("出现其他改造内容：" + AA)
                    break
            elif houseTransformType == "2":
                IA = "■"  # 否
            else:
                print("出现其他是否改造：" + AA)
                break
            # 是否用作经营
            manageHave = detail["data"]["manageHave"]
            if manageHave == "1":
                KA = "■"  # 是
                # 经营审批手续
                managePermit = detail["data"]["managePermit"]
                if managePermit == "1":
                    LA = "■"
                elif managePermit == "2":
                    LB = "■"
                else:
                    print("出现其他经营审批手续：" + AA)
                    break
                # 主要用途
                manageType = detail["data"]["manageType"]
                if len(manageType) > 1:
                    print("出现多次主要用途：" + AA)
                    break
                if manageType[0] == "1":
                    MA = "■"
                elif manageType[0] == "3":
                    MB = "■"
                elif manageType[0] == "4":
                    MC = "■"
                else:
                    print("出现其他主要用途：" + AA)
                    break
            elif manageHave == "2":
                KB = "■"  # 否
            else:
                print("出现其他是否用作经营：" + AA)
                break
            # 采暖用能
            heatEnergy = detail["data"]["heatEnergy"]
            if "1" in heatEnergy:
                NA = "■"
            if "2" in heatEnergy:
                NB = "■"
            if "3" in heatEnergy:
                NC = "■"
            if "4" in heatEnergy:
                ND = "■"
            if "5" in heatEnergy:
                NE = "■"
                NF = detail["data"]["heatEnergyOther"]
            # 炊事用能
            cookEnergy = detail["data"]["cookEnergy"]
            if "1" in cookEnergy:
                OA = "■"
            if "2" in cookEnergy:
                OB = "■"
            if "3" in cookEnergy:
                OC = "■"
            if "4" in cookEnergy:
                OD = "■"
            if "5" in cookEnergy:
                OE = "■"
                OF = detail["data"]["cookEnergyOther"]
            # 安全隐患初判
            houseSafeJudge = detail["data"]["houseSafeJudge"]
            if houseSafeJudge == "2":
                print("出现其他安全隐患初判：" + AA)

            # 遍历所有表格进行替换。
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                #print(run.text)
                                run.text = run.text.replace("AA", AA)
                                run.text = run.text.replace("AB", AB)
                                run.text = run.text.replace("BA", BA)
                                run.text = run.text.replace("BB", BB)
                                run.text = run.text.replace("BC", BC)
                                run.text = run.text.replace("BD", str(BD))
                                run.text = run.text.replace("CA", CA)
                                run.text = run.text.replace("CB", CB)
                                run.text = run.text.replace("CC", CC)
                                run.text = run.text.replace("CD", CD)
                                run.text = run.text.replace("CE", CE)
                                run.text = run.text.replace("DA", DA)
                                run.text = run.text.replace("DB", DB)
                                run.text = run.text.replace("DC", DC)
                                run.text = run.text.replace("DD", DD)
                                run.text = run.text.replace("DE", DE)
                                run.text = run.text.replace("EA", EA)
                                run.text = run.text.replace("EB", EB)
                                run.text = run.text.replace("EC", EC)
                                run.text = run.text.replace("ED", ED)
                                run.text = run.text.replace("EE", EE)
                                run.text = run.text.replace("FA", FA)
                                run.text = run.text.replace("FB", FB)
                                run.text = run.text.replace("GA", GA)
                                run.text = run.text.replace("GB", GB)
                                run.text = run.text.replace("GC", GC)
                                run.text = run.text.replace("GD", GD)
                                run.text = run.text.replace("HA", HA)
                                run.text = run.text.replace("HB", HB)
                                run.text = run.text.replace("HC", HC)
                                run.text = run.text.replace("HD", HD)
                                run.text = run.text.replace("IA", IA)
                                run.text = run.text.replace("IB", IB)
                                run.text = run.text.replace("IC", IC)
                                run.text = run.text.replace("ID", ID)
                                run.text = run.text.replace("JA", JA)
                                run.text = run.text.replace("JB", JB)
                                run.text = run.text.replace("JC", JC)
                                run.text = run.text.replace("JD", JD)
                                run.text = run.text.replace("JE", JE)
                                run.text = run.text.replace("JF", JF)
                                run.text = run.text.replace("KA", KA)
                                run.text = run.text.replace("KB", KB)
                                run.text = run.text.replace("LA", LA)
                                run.text = run.text.replace("LB", LB)
                                run.text = run.text.replace("MA", MA)
                                run.text = run.text.replace("MB", MB)
                                run.text = run.text.replace("MC", MC)
                                run.text = run.text.replace("NA", NA)
                                run.text = run.text.replace("NB", NB)
                                run.text = run.text.replace("NC", NC)
                                run.text = run.text.replace("ND", ND)
                                run.text = run.text.replace("NE", NE)
                                run.text = run.text.replace("NF", NF)
                                run.text = run.text.replace("OA", OA)
                                run.text = run.text.replace("OB", OB)
                                run.text = run.text.replace("OC", OC)
                                run.text = run.text.replace("OD", OD)
                                run.text = run.text.replace("OE", OE)
                                run.text = run.text.replace("OF", OF)

            doc.save(output_path + AA + "（" + detail["data"]["id"] + "）.docx")


def main():
    check_and_change()


if __name__ == '__main__':
    main()
